import io
import os
import re
import uvicorn
import tempfile
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, HTMLResponse
from pypdf import PdfReader
from docx import Document as DocxDocument
import easyocr
import whisper
import yt_dlp

# Configuración de FFmpeg
os.environ["PATH"] += os.pathsep + os.getcwd()

app = FastAPI(title="Sistema de Transcripción de Contenidos Educativos")

# CORS para permitir peticiones desde frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

print("=== SISTEMA DE TRANSCRIPCIÓN EDUCATIVA ===")

# Verificar FFmpeg
ffmpeg_available = False
if os.path.exists("ffmpeg.exe"):
    print("✅ FFmpeg detectado (ffmpeg.exe)")
    ffmpeg_available = True
elif os.system("ffmpeg -version > nul 2>&1" if os.name == 'nt' else "ffmpeg -version > /dev/null 2>&1") == 0:
    print("✅ FFmpeg detectado correctamente en PATH")
    ffmpeg_available = True
else:
    print("❌ ERROR CRÍTICO: FFmpeg NO encontrado")
    print("   SOLUCIÓN: Descarga FFmpeg de https://ffmpeg.org/download.html")
    print("   - Windows: Descarga ffmpeg.exe y colócalo en la misma carpeta")
    print("   - Linux/Mac: sudo apt install ffmpeg  o  brew install ffmpeg")

# Cargar motores
print("🔄 Cargando OCR (EasyOCR)...")
ocr_reader = easyocr.Reader(['es', 'en'], gpu=False)

print("🔄 Cargando Whisper (Transcripción de audio)...")
audio_model = whisper.load_model("base")

print("✅ TODOS LOS MOTORES LISTOS")
print("🌐 Servidor iniciado en http://localhost:7000\n")

# ==================== FUNCIONES DE EXTRACCIÓN ====================

def extract_pdf(file_bytes):
    """Extrae texto de PDFs"""
    try:
        reader_pdf = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader_pdf.pages:
            text += page.extract_text() + "\n"
        return text.strip() or "PDF procesado sin texto seleccionable."
    except Exception as e:
        return f"Error PDF: {str(e)}"

def extract_image_ocr(file_bytes):
    """OCR en imágenes"""
    try:
        return " ".join(ocr_reader.readtext(file_bytes, detail=0))
    except Exception as e:
        return f"Error OCR: {str(e)}"

def transcribe_audio_video(file_bytes, file_ext):
    """Transcripción de audio/video con timestamps"""
    try:
        with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        # Transcribir con timestamps
        result = audio_model.transcribe(tmp_path, fp16=False, verbose=False)
        
        os.remove(tmp_path)
        
        # Formatear con timestamps
        formatted_text = format_transcription_with_timestamps(result)
        
        return formatted_text
    except Exception as e:
        return f"Error Transcripción: {str(e)}"

def format_transcription_with_timestamps(result):
    """Formatea la transcripción con marcas de tiempo"""
    if "segments" not in result:
        return result.get("text", "")
    
    formatted = []
    for segment in result["segments"]:
        start_time = format_timestamp(segment["start"])
        text = segment["text"].strip()
        formatted.append(f"[{start_time}] {text}")
    
    return "\n".join(formatted)

def format_timestamp(seconds):
    """Convierte segundos a formato MM:SS"""
    mins = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{mins:02d}:{secs:02d}"

def transcribe_youtube(url):
    """Descarga audio de YouTube y transcribe - VERSIÓN MEJORADA"""
    audio_path = None
    
    # Verificar FFmpeg
    if not ffmpeg_available:
        return "ERROR: FFmpeg no está instalado. Es necesario para descargar audio de YouTube. Descárgalo de https://ffmpeg.org/download.html", None
    
    try:
        # Crear directorio temporal
        temp_dir = tempfile.gettempdir()
        temp_filename = os.path.join(temp_dir, f"yt_audio_{os.getpid()}")
        
        print(f"\n{'='*60}")
        print(f"🎬 INICIANDO DESCARGA DE YOUTUBE")
        print(f"{'='*60}")
        print(f"📍 URL: {url}")
        print(f"📁 Archivo temporal: {temp_filename}")
        
        # Configuración mejorada de yt-dlp con más opciones
        ydl_opts = {
            'format': 'bestaudio/best',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            'outtmpl': temp_filename,
            'quiet': True,
            'no_warnings': False,
            'verbose': False, 
            'nocheckcertificate': True,
            'ignoreerrors': False,
            'no_color': True,
            'extractaudio': True,
            'socket_timeout': 30,
            'retries': 3,
            'fragment_retries': 3,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-us,en;q=0.5',
                'Sec-Fetch-Mode': 'navigate',
            }
        }
        
        print("🔄 Extrayendo información del video...")
        
        # Descargar audio
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            try:
                info = ydl.extract_info(url, download=True)
                title = info.get('title', 'Video de YouTube')
                duration = info.get('duration', 0)
                print(f"✅ Video encontrado: {title}")
                print(f"⏱️  Duración: {duration} segundos")
            except yt_dlp.utils.DownloadError as e:
                error_msg = str(e)
                print(f"❌ Error de descarga: {error_msg}")
                
                if "Video unavailable" in error_msg:
                    return "❌ El video no está disponible. Puede ser privado, estar bloqueado en tu región o haber sido eliminado.", None
                elif "Sign in" in error_msg or "age" in error_msg.lower():
                    return "❌ Este video requiere iniciar sesión o verificación de edad. No se puede descargar.", None
                elif "Copyright" in error_msg:
                    return "❌ Este video tiene restricciones de copyright y no se puede descargar.", None
                else:
                    return f"❌ Error al descargar: {error_msg}", None
            except Exception as e:
                print(f"❌ Error inesperado en descarga: {str(e)}")
                return f"❌ Error al procesar el video: {str(e)}", None
        
        # Buscar el archivo de audio generado
        print("🔍 Buscando archivo de audio descargado...")
        audio_path = f"{temp_filename}.mp3"
        
        if not os.path.exists(audio_path):
            for ext in ['.m4a', '.webm', '.opus', '.ogg']:
                alt_path = f"{temp_filename}{ext}"
                if os.path.exists(alt_path):
                    audio_path = alt_path
                    print(f"✅ Archivo encontrado: {alt_path}")
                    break
        
        if not os.path.exists(audio_path):
            print("❌ No se encontró el archivo de audio después de la descarga")
            print(f"   Buscado: {audio_path}")
            print(f"   Archivos en {temp_dir}:")
            for f in os.listdir(temp_dir):
                if 'yt_audio' in f:
                    print(f"   - {f}")
            return "❌ Error: El audio se descargó pero no se encontró el archivo. Verifica que FFmpeg esté correctamente instalado.", None
        
        print(f"✅ Archivo de audio listo: {os.path.basename(audio_path)}")
        file_size = os.path.getsize(audio_path) / (1024 * 1024)  # MB
        print(f"📦 Tamaño: {file_size:.2f} MB")
        
        print("🔄 Iniciando transcripción con Whisper...")
        
        # Transcribir
        result = audio_model.transcribe(audio_path, fp16=False, verbose=False)
        formatted_text = format_transcription_with_timestamps(result)
        
        print(f"✅ Transcripción completada: {len(formatted_text)} caracteres")
        print(f"{'='*60}\n")
        
        return formatted_text, title
        
    except Exception as e:
        error_msg = f"❌ Error inesperado: {str(e)}"
        print(error_msg)
        import traceback
        print(traceback.format_exc())
        return error_msg, None
    
    finally:
        # Limpiar archivo temporal
        if audio_path and os.path.exists(audio_path):
            try:
                os.remove(audio_path)
                print(f"🗑️  Archivo temporal eliminado: {os.path.basename(audio_path)}")
            except Exception as e:
                print(f"⚠️  No se pudo eliminar archivo temporal: {e}")

def extract_plain_text(file_bytes):
    """Lee archivos de texto plano"""
    return file_bytes.decode("utf-8", errors="ignore")

# ==================== GENERACIÓN DE ARCHIVOS ====================

def generate_docx(text, filename="transcripcion.docx"):
    """Genera un archivo Word con el texto"""
    doc = DocxDocument()
    doc.add_heading('Transcripción de Contenido Educativo', 0)
    
    # Agregar párrafos
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    
    # Guardar en memoria
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.getvalue()

# ==================== HTML EMBEBIDO ====================

HTML_CONTENT = """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Sistema de Transcripción Educativa</title>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <style>
        :root {
            --primary: #0891b2;
            --primary-hover: #0e7490;
            --secondary: #06b6d4;
            --accent: #10b981;
            --bg-body: #f0fdfa;
            --bg-card: #ffffff;
            --bg-card-hover: #ecfeff;
            --text-main: #0f172a;
            --text-sub: #475569;
            --border: #cbd5e1;
            --radius: 16px;
            --shadow: 0 4px 6px -1px rgba(8, 145, 178, 0.1), 0 2px 4px -1px rgba(8, 145, 178, 0.06);
            --shadow-lg: 0 20px 25px -5px rgba(8, 145, 178, 0.1), 0 10px 10px -5px rgba(8, 145, 178, 0.04);
        }
        
        * { box-sizing: border-box; margin: 0; padding: 0; }
        
        body {
            font-family: 'Inter', system-ui, -apple-system, sans-serif;
            background: linear-gradient(135deg, #ecfeff 0%, #f0fdfa 100%);
            color: var(--text-main);
            min-height: 100vh;
            padding: 20px;
        }
        
        .container {
            max-width: 1200px;
            margin: 0 auto;
        }
        
        .header {
            text-align: center;
            margin-bottom: 40px;
            padding: 30px;
            background: white;
            border-radius: var(--radius);
            box-shadow: var(--shadow);
        }
        
        .header h1 {
            font-size: 2.5rem;
            font-weight: 700;
            background: linear-gradient(135deg, var(--primary), var(--accent));
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 10px;
        }
        
        .header p {
            color: var(--text-sub);
            font-size: 1.1rem;
        }
        
        .grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 30px;
            margin-bottom: 30px;
        }
        
        @media (max-width: 768px) {
            .grid { grid-template-columns: 1fr; }
        }
        
        .card {
            background: var(--bg-card);
            border-radius: var(--radius);
            padding: 30px;
            box-shadow: var(--shadow);
            transition: all 0.3s ease;
        }
        
        .card:hover {
            box-shadow: var(--shadow-lg);
            transform: translateY(-4px);
        }
        
        .card-icon {
            width: 60px;
            height: 60px;
            background: linear-gradient(135deg, var(--primary), var(--accent));
            border-radius: 12px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 24px;
            color: white;
            margin-bottom: 20px;
        }
        
        .card h3 {
            font-size: 1.5rem;
            margin-bottom: 10px;
            color: var(--text-main);
        }
        
        .card p {
            color: var(--text-sub);
            margin-bottom: 20px;
            line-height: 1.6;
        }
        
        .upload-zone {
            border: 3px dashed var(--border);
            border-radius: var(--radius);
            padding: 50px 20px;
            text-align: center;
            cursor: pointer;
            transition: all 0.3s ease;
            background: var(--bg-card-hover);
        }
        
        .upload-zone:hover {
            border-color: var(--primary);
            background: white;
        }
        
        .upload-zone i {
            font-size: 48px;
            color: var(--primary);
            margin-bottom: 15px;
        }
        
        .upload-zone p {
            color: var(--text-sub);
            font-size: 1rem;
        }
        
        .youtube-input {
            display: flex;
            gap: 10px;
        }
        
        input[type="text"] {
            flex: 1;
            padding: 15px;
            border: 2px solid var(--border);
            border-radius: 12px;
            font-size: 1rem;
            transition: border-color 0.3s;
        }
        
        input[type="text"]:focus {
            outline: none;
            border-color: var(--primary);
        }
        
        .btn {
            padding: 15px 30px;
            border: none;
            border-radius: 12px;
            font-size: 1rem;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
            display: inline-flex;
            align-items: center;
            gap: 8px;
        }
        
        .btn-primary {
            background: linear-gradient(135deg, var(--primary), var(--accent));
            color: white;
        }
        
        .btn-primary:hover {
            transform: translateY(-2px);
            box-shadow: var(--shadow-lg);
        }
        
        .btn-primary:disabled {
            opacity: 0.5;
            cursor: not-allowed;
            transform: none;
        }
        
        .btn-outline {
            background: transparent;
            border: 2px solid var(--primary);
            color: var(--primary);
        }
        
        .btn-outline:hover {
            background: var(--primary);
            color: white;
        }
        
        .btn-small {
            padding: 10px 20px;
            font-size: 0.9rem;
        }
        
        .loading, .results {
            display: none;
        }
        
        .loading.active, .results.active {
            display: block;
        }
        
        .loading {
            text-align: center;
            padding: 60px 30px;
            background: white;
            border-radius: var(--radius);
            box-shadow: var(--shadow);
        }
        
        .spinner {
            width: 60px;
            height: 60px;
            border: 5px solid var(--bg-card-hover);
            border-top-color: var(--primary);
            border-radius: 50%;
            animation: spin 1s linear infinite;
            margin: 0 auto 20px;
        }
        
        @keyframes spin {
            to { transform: rotate(360deg); }
        }
        
        .loading h3 {
            color: var(--text-main);
            margin-bottom: 10px;
        }
        
        .loading p {
            color: var(--text-sub);
        }
        
        .results {
            background: white;
            border-radius: var(--radius);
            padding: 30px;
            box-shadow: var(--shadow);
        }
        
        .result-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 20px;
            padding-bottom: 20px;
            border-bottom: 2px solid var(--bg-card-hover);
        }
        
        .result-info h3 {
            font-size: 1.5rem;
            color: var(--text-main);
            margin-bottom: 5px;
        }
        
        .result-info p {
            color: var(--text-sub);
        }
        
        .badge {
            padding: 8px 16px;
            background: linear-gradient(135deg, var(--primary), var(--accent));
            color: white;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
        }
        
        .stats {
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 15px;
            margin-bottom: 20px;
        }
        
        .stat-card {
            background: var(--bg-card-hover);
            padding: 15px;
            border-radius: 12px;
            text-align: center;
        }
        
        .stat-card i {
            font-size: 24px;
            color: var(--primary);
            margin-bottom: 8px;
        }
        
        .stat-card strong {
            display: block;
            font-size: 1.5rem;
            color: var(--text-main);
            margin-bottom: 5px;
        }
        
        .stat-card span {
            color: var(--text-sub);
            font-size: 0.9rem;
        }
        
        textarea {
            width: 100%;
            min-height: 300px;
            padding: 20px;
            border: 2px solid var(--border);
            border-radius: 12px;
            font-family: 'Courier New', monospace;
            font-size: 0.95rem;
            resize: vertical;
            margin-bottom: 20px;
        }
        
        textarea:focus {
            outline: none;
            border-color: var(--primary);
        }
        
        .actions {
            display: flex;
            gap: 10px;
            flex-wrap: wrap;
        }
        
        input[type="file"] {
            display: none;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1><i class="fa-solid fa-microphone-lines"></i> Sistema de Transcripción</h1>
            <p>Convierte audio, video, imágenes y documentos a texto</p>
        </div>
        
        <div class="grid">
            <div class="card">
                <div class="card-icon">
                    <i class="fa-solid fa-file-arrow-up"></i>
                </div>
                <h3>Subir Archivo</h3>
                <p>Soporta: MP3, MP4, PDF, imágenes, texto</p>
                <input type="file" id="file-input" accept=".mp3,.mp4,.wav,.m4a,.png,.jpg,.jpeg,.pdf,.txt,.md">
                <div class="upload-zone" id="upload-zone">
                    <i class="fa-solid fa-cloud-arrow-up"></i>
                    <p><strong>Haz clic o arrastra un archivo aquí</strong></p>
                    <p>Máximo 500MB</p>
                </div>
            </div>
            
            <div class="card">
                <div class="card-icon">
                    <i class="fa-brands fa-youtube"></i>
                </div>
                <h3>YouTube</h3>
                <p>Transcribe directamente desde YouTube</p>
                <div class="youtube-input">
                    <input type="text" id="youtube-url" placeholder="Pega la URL del video aquí...">
                    <button class="btn btn-primary" onclick="transcribeYouTube()">
                        <i class="fa-solid fa-play"></i> Transcribir
                    </button>
                </div>
            </div>
        </div>
        
        <div class="loading" id="loading">
            <div class="spinner"></div>
            <h3>Procesando...</h3>
            <p>Esto puede tomar unos momentos</p>
        </div>
        
        <div class="results" id="results">
            <div class="result-header">
                <div class="result-info">
                    <h3 id="result-filename">nombre_archivo.mp3</h3>
                    <p id="result-processor">Procesador utilizado</p>
                </div>
                <span class="badge" id="result-badge">COMPLETADO</span>
            </div>
            
            <div class="stats">
                <div class="stat-card">
                    <i class="fa-solid fa-font"></i>
                    <strong id="stat-chars">0</strong>
                    <span>Caracteres</span>
                </div>
                <div class="stat-card">
                    <i class="fa-solid fa-spell-check"></i>
                    <strong id="stat-words">0</strong>
                    <span>Palabras</span>
                </div>
                <div class="stat-card">
                    <i class="fa-solid fa-list-ol"></i>
                    <strong id="stat-lines">0</strong>
                    <span>Líneas</span>
                </div>
            </div>
            
            <textarea id="result-text" placeholder="El texto transcrito aparecerá aquí..."></textarea>
            
            <div class="actions">
                <button class="btn btn-primary btn-small" onclick="copyText()">
                    <i class="fa-solid fa-copy"></i> Copiar
                </button>
                <button class="btn btn-primary btn-small" onclick="downloadFile('txt')">
                    <i class="fa-solid fa-download"></i> TXT
                </button>
                <button class="btn btn-primary btn-small" onclick="downloadFile('docx')">
                    <i class="fa-solid fa-file-word"></i> DOCX
                </button>
                <button class="btn btn-outline btn-small" onclick="resetApp()">
                    <i class="fa-solid fa-rotate-left"></i> Nuevo
                </button>
            </div>
        </div>
    </div>
    
    <script>
        const fileInput = document.getElementById('file-input');
        const uploadZone = document.getElementById('upload-zone');
        const loading = document.getElementById('loading');
        const results = document.getElementById('results');
        const resultText = document.getElementById('result-text');
        
        uploadZone.addEventListener('click', () => fileInput.click());
        uploadZone.addEventListener('dragover', (e) => {
            e.preventDefault();
            uploadZone.style.borderColor = 'var(--primary)';
            uploadZone.style.background = 'white';
        });
        uploadZone.addEventListener('dragleave', () => {
            uploadZone.style.borderColor = 'var(--border)';
            uploadZone.style.background = 'var(--bg-card-hover)';
        });
        uploadZone.addEventListener('drop', (e) => {
            e.preventDefault();
            uploadZone.style.borderColor = 'var(--border)';
            uploadZone.style.background = 'var(--bg-card-hover)';
            if (e.dataTransfer.files.length) processFile(e.dataTransfer.files[0]);
        });
        fileInput.addEventListener('change', (e) => {
            if (e.target.files.length) processFile(e.target.files[0]);
        });
        
        async function processFile(file) {
            showLoading();
            const formData = new FormData();
            formData.append("file", file);
            
            try {
                const response = await fetch("/upload/", {
                    method: "POST",
                    body: formData
                });
                const data = await response.json();
                if (!response.ok) throw new Error(data.detail || "Error en el servidor");
                showResults(data, file.name);
            } catch (error) {
                alert("Error: " + error.message);
                hideLoading();
            }
        }
        
        async function transcribeYouTube() {
            const url = document.getElementById('youtube-url').value.trim();
            if (!url) {
                alert("Por favor ingresa una URL");
                return;
            }
            
            showLoading();
            const formData = new FormData();
            formData.append("url", url);
            
            try {
                const response = await fetch("/transcribe-youtube/", {
                    method: "POST",
                    body: formData
                });
                const data = await response.json();
                if (data.error) throw new Error(data.message);
                showResults(data, data.filename);
            } catch (error) {
                alert("Error: " + error.message);
                hideLoading();
            }
        }
        
        function showResults(data, filename) {
            hideLoading();
            results.classList.add('active');
            
            document.getElementById('result-filename').textContent = filename;
            document.getElementById('result-processor').textContent = data.processor;
            document.getElementById('result-badge').textContent = 'COMPLETADO';
            resultText.value = data.extracted_data;
            
            const text = data.extracted_data;
            document.getElementById('stat-chars').textContent = text.length.toLocaleString();
            document.getElementById('stat-words').textContent = text.split(/\\s+/).filter(w => w).length.toLocaleString();
            document.getElementById('stat-lines').textContent = text.split('\\n').length.toLocaleString();
        }
        
        function showLoading() {
            results.classList.remove('active');
            loading.classList.add('active');
        }
        
        function hideLoading() {
            loading.classList.remove('active');
        }
        
        function copyText() {
            resultText.select();
            navigator.clipboard.writeText(resultText.value);
            alert("✅ Texto copiado al portapapeles");
        }
        
        async function downloadFile(format) {
            const text = resultText.value;
            const formData = new FormData();
            formData.append("text", text);
            formData.append("format", format);
            
            try {
                const response = await fetch("/download/", {
                    method: "POST",
                    body: formData
                });
                const blob = await response.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = `transcripcion.${format}`;
                a.click();
            } catch (error) {
                alert("Error al descargar: " + error.message);
            }
        }
        
        function resetApp() {
            fileInput.value = '';
            document.getElementById('youtube-url').value = '';
            resultText.value = '';
            results.classList.remove('active');
        }
    </script>
</body>
</html>
"""

# ==================== ENDPOINTS ====================

@app.get("/", response_class=HTMLResponse)
def root():
    """Sirve la interfaz HTML"""
    return HTML_CONTENT

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    """Procesa archivos subidos"""
    content = await file.read()
    filename = file.filename.lower()
    
    # IMÁGENES → OCR
    if filename.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp")):
        data = extract_image_ocr(content)
        proc = "OCR (EasyOCR)"
        
    # AUDIO/VIDEO → TRANSCRIPCIÓN
    elif filename.endswith((".mp3", ".wav", ".mp4", ".mpeg", ".m4a", ".ogg", ".webm")):
        ext = os.path.splitext(filename)[1]
        data = transcribe_audio_video(content, ext)
        proc = "Transcripción (Whisper AI)"
        
    # PDF
    elif filename.endswith(".pdf"):
        data = extract_pdf(content)
        proc = "Extracción PDF"
        
    # TEXTO PLANO
    elif filename.endswith((".txt", ".md", ".py", ".json", ".csv")):
        data = extract_plain_text(content)
        proc = "Texto Plano"
        
    else:
        data = "Formato no soportado."
        proc = "Desconocido"

    return {
        "filename": file.filename,
        "processor": proc,
        "extracted_data": data,
        "char_count": len(data),
        "word_count": len(data.split())
    }

@app.post("/transcribe-youtube/")
async def transcribe_youtube_endpoint(url: str = Form(...)):
    """Transcribe videos de YouTube"""
    text, title = transcribe_youtube(url)
    
    if title is None:
        return {
            "error": True,
            "message": text
        }
    
    return {
        "filename": title,
        "processor": "YouTube + Whisper AI",
        "extracted_data": text,
        "char_count": len(text),
        "word_count": len(text.split())
    }

@app.post("/download/")
async def download_text(text: str = Form(...), format: str = Form("txt")):
    """Genera archivo descargable (TXT o DOCX)"""
    if format == "docx":
        content = generate_docx(text)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "transcripcion.docx"
    else:
        content = text.encode("utf-8")
        media_type = "text/plain"
        filename = "transcripcion.txt"
    
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

if __name__ == "__main__":
    # Abrir navegador automáticamente
    import webbrowser
    from threading import Timer
    
    def open_browser():
        webbrowser.open("http://localhost:7000")
    
    Timer(1.5, open_browser).start()
    
    uvicorn.run(app, host="0.0.0.0", port=7000)
