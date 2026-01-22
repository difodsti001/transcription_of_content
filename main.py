import io
import os
import re
import uvicorn
import tempfile
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import Response
from pypdf import PdfReader
from docx import Document as DocxDocument
import easyocr
import whisper
import yt_dlp

# Configuración de FFmpeg
os.environ["PATH"] += os.pathsep + os.getcwd()

app = FastAPI(title="Sistema de Transcripción de Contenidos Educativos")

# CORS para permitir peticiones desde frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="static"), name="static")

print("=== SISTEMA DE TRANSCRIPCIÓN EDUCATIVA ===")

# Verificar FFmpeg
if os.path.exists("ffmpeg.exe") or os.system("ffmpeg -version > nul 2>&1") == 0:
    print("✅ FFmpeg detectado correctamente")
else:
    print("⚠️ ADVERTENCIA: FFmpeg no encontrado")

# Cargar motores
print("🔄 Cargando OCR (EasyOCR)...")
ocr_reader = easyocr.Reader(['es', 'en'], gpu=False)

print("🔄 Cargando Whisper (Transcripción de audio)...")
audio_model = whisper.load_model("base")

print("✅ TODOS LOS MOTORES LISTOS\n")

# ==================== FUNCIONES DE EXTRACCIÓN ====================

def extract_pdf(file_bytes):
    """Extrae texto de PDFs"""
    try:
        reader_pdf = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader_pdf.pages:
            text += page.extract_text() + "\n"
        return text.strip() or "PDF procesado sin texto seleccionable."
    except Exception as e:
        return f"Error PDF: {str(e)}"

def extract_image_ocr(file_bytes):
    """OCR en imágenes"""
    try:
        return " ".join(ocr_reader.readtext(file_bytes, detail=0))
    except Exception as e:
        return f"Error OCR: {str(e)}"

def transcribe_audio_video(file_bytes, file_ext):
    """Transcripción de audio/video con timestamps"""
    try:
        with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        # Transcribir con timestamps
        result = audio_model.transcribe(tmp_path, fp16=False, verbose=False)
        
        os.remove(tmp_path)
        
        # Formatear con timestamps
        formatted_text = format_transcription_with_timestamps(result)
        
        return formatted_text
    except Exception as e:
        return f"Error Transcripción: {str(e)}"

def format_transcription_with_timestamps(result):
    """Formatea la transcripción con marcas de tiempo"""
    if "segments" not in result:
        return result.get("text", "")
    
    formatted = []
    for segment in result["segments"]:
        start_time = format_timestamp(segment["start"])
        text = segment["text"].strip()
        formatted.append(f"[{start_time}] {text}")
    
    return "\n".join(formatted)

def format_timestamp(seconds):
    """Convierte segundos a formato MM:SS"""
    mins = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{mins:02d}:{secs:02d}"

def transcribe_youtube(url):
    """Descarga audio de YouTube y transcribe"""
    try:
        # Configuración de yt-dlp
        ydl_opts = {
            'format': 'bestaudio/best',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            'outtmpl': 'temp_yt_audio.%(ext)s',
            'quiet': True,
            'no_warnings': True,
        }
        
        # Descargar audio
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            title = info.get('title', 'Video de YouTube')
        
        # Transcribir
        audio_path = "temp_yt_audio.mp3"
        if not os.path.exists(audio_path):
            return "Error: No se pudo descargar el audio", None
        
        result = audio_model.transcribe(audio_path, fp16=False, verbose=False)
        formatted_text = format_transcription_with_timestamps(result)
        
        # Limpiar archivo temporal
        os.remove(audio_path)
        
        return formatted_text, title
        
    except Exception as e:
        return f"Error YouTube: {str(e)}", None

def extract_plain_text(file_bytes):
    """Lee archivos de texto plano"""
    return file_bytes.decode("utf-8", errors="ignore")

# ==================== GENERACIÓN DE ARCHIVOS ====================

def generate_docx(text, filename="transcripcion.docx"):
    """Genera un archivo Word con el texto"""
    doc = DocxDocument()
    doc.add_heading('Transcripción de Contenido Educativo', 0)
    
    # Agregar párrafos
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    
    # Guardar en memoria
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.getvalue()

# ==================== ENDPOINTS ====================

@app.get("/", response_class= HTMLResponse)
def root():
    with open("static/index.html", "r", encoding="utf-8") as f:
        return f.read()

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    """Procesa archivos subidos"""
    content = await file.read()
    filename = file.filename.lower()
    
    # IMÁGENES → OCR
    if filename.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp")):
        data = extract_image_ocr(content)
        proc = "OCR (EasyOCR)"
        
    # AUDIO/VIDEO → TRANSCRIPCIÓN
    elif filename.endswith((".mp3", ".wav", ".mp4", ".mpeg", ".m4a", ".ogg", ".webm")):
        ext = os.path.splitext(filename)[1]
        data = transcribe_audio_video(content, ext)
        proc = "Transcripción (Whisper AI)"
        
    # PDF
    elif filename.endswith(".pdf"):
        data = extract_pdf(content)
        proc = "Extracción PDF"
        
    # TEXTO PLANO
    elif filename.endswith((".txt", ".md", ".py", ".json", ".csv")):
        data = extract_plain_text(content)
        proc = "Texto Plano"
        
    else:
        data = "Formato no soportado."
        proc = "Desconocido"

    return {
        "filename": file.filename,
        "processor": proc,
        "extracted_data": data,
        "char_count": len(data),
        "word_count": len(data.split())
    }

@app.post("/transcribe-youtube/")
async def transcribe_youtube_endpoint(url: str = Form(...)):
    """Transcribe videos de YouTube"""
    text, title = transcribe_youtube(url)
    
    if title is None:
        return {
            "error": True,
            "message": text
        }
    
    return {
        "filename": title,
        "processor": "YouTube + Whisper AI",
        "extracted_data": text,
        "char_count": len(text),
        "word_count": len(text.split())
    }

@app.post("/download/")
async def download_text(text: str = Form(...), format: str = Form("txt")):
    """Genera archivo descargable (TXT o DOCX)"""
    if format == "docx":
        content = generate_docx(text)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "transcripcion.docx"
    else:
        content = text.encode("utf-8")
        media_type = "text/plain"
        filename = "transcripcion.txt"
    
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=7000)